#!/usr/bin/env python
# coding: utf-8


# Inputs and Outputs
# Custom list of UK and US Words are zipped together 

import re
import docx2txt
import timeit
from collections import Counter
from docx import Document
from tqdm import tqdm

my_text = docx2txt.process(r"C:\Users\91824\Downloads\Finding-US-UK-Words\Project_Transform\Coal Ash Review R2-v5.docx")
filename = r"C:\Users\91824\Downloads\Finding-US-UK-Words\Project_Transform\Coal Ash Review R2-v5.docx"
output_name = r"C:\Users\91824\Downloads\Finding-US-UK-Words\Project_Transform\Coal Ash Review R2-v5_final.docx"
doc = Document(filename)


def Initial():
    start_time = timeit.default_timer()
#     my_input = input("Please enter an input: ")

    
    with open(r"C:\Users\91824\Downloads\Finding-US-UK-Words\Project_Transform\uk.txt") as f1:
        x= f1.read().split()

    with open(r"C:\Users\91824\Downloads\Finding-US-UK-Words\Project_Transform\Us.txt") as f2:
        y= f2.read().split()

    US_UK  = dict(zip(x, y))
    return US_UK,x,y,start_time


# There are three possibities for a word to get matched with either British or American words
# Total Matched Words(English or British) = Uppercase Words + Lowercase Words + Title Words


def Word_count():
    US_UK,x,y,start_time = Initial()
    Title_word_US = list(map(str.title,y))
    Upper_word_US = list(map(str.upper,y))
    Lower_word_US = list(map(str.lower,y))

    Title_word_UK = list(map(str.title,x))
    Upper_word_UK = list(map(str.upper,x))
    Lower_word_UK = list(map(str.lower,x))

    Us_words = [Title_word_US,Upper_word_US,Lower_word_US]
    Uk_words = [Title_word_UK,Upper_word_UK,Lower_word_UK]


    Words_count_US = []
    for i in Us_words:
        vocab = i
        r = re.compile("|".join(r'\b%s?\b' % w for w in vocab))
        wordcount_us = Counter(re.findall(r, my_text))
# print(wordcount_us)
        if '/' or '' or 'to' or 'Dis' or 'To'or 'dis'in wordcount_us:
            del wordcount_us['/'],wordcount_us[''],wordcount_us['To'],wordcount_us['to'],wordcount_us['Dis'],wordcount_us['dis']
#         print(wordcount_us)
            Words_count_US.append(wordcount_us)


    Words_count_UK = []
    for i in Uk_words:
        vocab = i
        r = re.compile("|".join(r'\b%s?\b' % w for w in vocab))
        wordcount_uk = Counter(re.findall(r, my_text))
# print(wordcount_uk)
#     for k,v in (dict(wordcount_uk)).items:
#         print(k)
        if '/' or '' or 'to' or 'Dis' or 'To'or 'dis'in wordcount_uk:
            del wordcount_uk['/'],wordcount_uk[''],wordcount_uk['To'],wordcount_uk['to'],wordcount_uk['Dis'],wordcount_uk['dis']
            Words_count_UK.append(wordcount_uk)
        
    count_us = []
    for i in Words_count_US:
#     for y in i.values:
        count_us.append(sum(i.values()))

    count_uk = []
    for i in Words_count_UK:
#     for y in i.values:
        count_uk.append(sum(i.values()))
    

    return US_UK , count_us ,count_uk,Words_count_US,Words_count_UK

def Dict_words():
    US_UK , count_us ,count_uk,Words_count_US,Words_count_UK = Word_count()
    dict_uk = {}
    dict_us = {}

    dict_lower = dict((k.lower(), v.lower()) for k,v in US_UK.items())
    dict_upper = dict((k.upper(), v.upper()) for k,v in US_UK.items())
    dict_title = dict((k.title(), v.title()) for k,v in US_UK.items())

#Inverting Dict from US, UK as Key and Value

    inv_dict_lower = {v: k for k, v in dict_lower.items()}
    inv_dict_upper = {v: k for k, v in dict_upper.items()}
    inv_dict_title = {v: k for k, v in dict_title.items()}



    if count_us < count_uk:
        for i in Words_count_UK:
#     print(len(dict(i).values()))
            if len(dict(i).values()) > 0:
                for i in dict(i):
                    if i.istitle():
#                     print(str(i)+ ' : '+ str(dict_title.get(i)))
                        dict_us[i] = dict_title.get(i)
#                 print(dict_title.get(i))
                    if i.islower():
#                 print(str(i)+ ' : '+ str(dict_lower.get(i)))
                        dict_us[i] = dict_lower.get(i)
                    if i.isupper():
#                 print(str(i)+ ' : '+ str(dict_upper.get(i)))
                        dict_us[i] = dict_lower.get(i)
 
# print("----------------------------------------")
    if count_us > count_uk:
        for i in Words_count_US:
#     print(len(dict(i).values()))
            if len(dict(i).values()) > 0:
                for i in dict(i):
                    if i.istitle():
#                 print(str(i)+ ' : '+ str(inv_dict_title.get(i)))
                        dict_uk[i] = inv_dict_title.get(i)
#                 print(dict_title.get(i))
                    if i.islower():
#                 print(str(i)+ ' : '+ str(inv_dict_lower.get(i)))
                        dict_uk[i] = inv_dict_lower.get(i)
                    if i.isupper():
#                 print(str(i)+ ' : '+ str(inv_dict_upper.get(i)))
                        dict_uk[i] = inv_dict_upper.get(i)
    
    output = [dict_uk,dict_us]
    return output

def docx_replace_regex(doc_obj, regex , replace):

    for p in doc_obj.paragraphs:
        if regex.search(p.text):
            inline = p.runs
            # Loop added to work with runs (strings with same style)
            for i in range(len(inline)):
                if regex.search(inline[i].text):
                    text = regex.sub(replace, inline[i].text)
                    inline[i].text = text

    for table in doc_obj.tables:
        for row in table.rows:
            for cell in row.cells:
                docx_replace_regex(cell, regex , replace)

def converter():
    coute = []
    output = Dict_words()


    for i in output:
        if len(i) > 0:
#         print(range(i))
#      for i in tqdm(i, len(i)):
#         print(i)
            print("Processing...")
            print("                               ")
            for word, replacement in tqdm(i.items()):
                word_re=re.compile(r'\b%s\b' %word)
                docx_replace_regex(doc, word_re , replacement)
            print("                               ")
            
            doc.save(output_name)

    # for i in output:
    #     l = len(i)
    #     coute.append(l)

    # if coute[0] > coute[1]:
    #     print("    UK is dominating by %s " % str(round(sum(count_uk)/(sum(count_us)+sum(count_uk)),2)*100))
    #     print("-------------------------------")
    #     print("    Done with Replacing")
    # if coute[0] < coute[1]:
    #     print("    US is dominating by %s" % str(round(sum(count_us)/(sum(count_us)+sum(count_uk)),2)*100))
    #     print("-------------------------------")
    #     print("    Done with Replacing")
            

if __name__ == "__main__":
    start_time = timeit.default_timer()
    print("                               ")
    print("-------------------------------")
    print("   Developed by Makesh Krishna")
    print("-------------------------------")
    
    print("   Initialization has begun...")
    Initial()
    print("-------------------------------")
    US_UK , count_us ,count_uk,Words_count_US,Words_count_UK = Word_count()
    print("    US Word count: " + str(sum(count_us)))
    print("    UK Word count: " + str(sum(count_uk)))
    print("                               ")
    Dict_words()
    start_time_ = timeit.default_timer()
    converter()
    print("----------------------------------------")
    print("   Overall code Execution time : {0:.2f}".format(timeit.default_timer() - start_time))
    print("   Coverter code Execution time : {0:.2f}".format(timeit.default_timer() - start_time_))